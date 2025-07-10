#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import json
import os
import shutil
import sys
import threading
import tkinter as tk
import zipfile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from pathlib import Path
from tkinter import ttk, filedialog, messagebox

class I18n:
    """Internationalization support class;"""
    
    def __init__(self, default_locale='en'):
        self.current_locale = default_locale
        self.translations = {}
        self.load_translations()
    
    def load_translations(self):
        """Load translation files;"""
        locales_dir = Path(__file__).parent / 'locales'
        if not locales_dir.exists():
            return
        
        for locale_file in locales_dir.glob('*.json'):
            locale_name = locale_file.stem
            try:
                with open(locale_file, 'r', encoding='utf-8') as f:
                    self.translations[locale_name] = json.load(f)
            except Exception as e:
                print(f"Error loading translation file {locale_file}: {e}")
    
    def set_locale(self, locale):
        """Set current locale;"""
        if locale in self.translations:
            self.current_locale = locale
            return True
        return False
    
    def get_available_locales(self):
        """Get list of available locales;"""
        return list(self.translations.keys())
    
    def get_locale_display_name(self, locale):
        """Get user-friendly display name for locale;"""
        display_names = {
            'en': 'English',
            'zh_Hans': '简体中文',
            'zh_Hant': '繁體中文'
        }
        return display_names.get(locale, locale)
    
    def get_locale_from_display_name(self, display_name):
        """Get locale code from display name;"""
        display_names = {
            'English': 'en',
            '简体中文': 'zh_Hans',
            '繁體中文': 'zh_Hant'
        }
        return display_names.get(display_name, display_name)
    
    def t(self, key, **kwargs):
        """Get translated text;"""
        if self.current_locale not in self.translations:
            return key
        
        text = self.translations[self.current_locale].get(key, key)
        
        # Format with provided kwargs;
        if kwargs:
            try:
                return text.format(**kwargs)
            except (KeyError, ValueError):
                return text
        
        return text


class DuplicateDetectorApp:
    """Main application class for Word duplicate detector;"""
    
    def __init__(self, root):
        self.root = root
        self.i18n = I18n()
        
        # Try to detect system locale;
        try:
            import locale
            system_locale = locale.getdefaultlocale()[0]
            if system_locale:
                if system_locale.startswith('zh_Hans') or system_locale.startswith('zh_Hans'):
                    self.i18n.set_locale('zh_Hans')
                elif system_locale.startswith('zh_Hant') or system_locale.startswith('zh_Hant'):
                    self.i18n.set_locale('zh_Hant')
                elif system_locale.startswith('zh'):
                    # Default to simplified Chinese for other Chinese locales;
                    self.i18n.set_locale('zh_Hans')
        except:
            pass
        
        self.setup_window()
        self.setup_ui()
        
    def setup_window(self):
        """Setup main window properties;"""
        self.root.title(self.i18n.t('window_title'))
        self.root.geometry("700x650")
        self.root.resizable(True, True)
        
        # Set window icon if available;
        try:
            # Add icon file handling here if needed;
            pass
        except:
            pass
    
    def setup_ui(self):
        """Setup user interface;"""
        # Main frame;
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Configure grid weights;
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        
        # Title;
        title_label = ttk.Label(main_frame, text=self.i18n.t('app_title'), 
                               font=('Arial', 16, 'bold'))
        title_label.grid(row=0, column=0, columnspan=3, pady=(0, 5))
        
        # Subtitle;
        subtitle_label = ttk.Label(main_frame, text=self.i18n.t('subtitle'), 
                                  font=('Arial', 9), foreground='#666666')
        subtitle_label.grid(row=1, column=0, columnspan=3, pady=(0, 20))
        
        # Language selection;
        lang_frame = ttk.Frame(main_frame)
        lang_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        ttk.Label(lang_frame, text=self.i18n.t('language') + ":").pack(side=tk.LEFT)
        self.lang_var = tk.StringVar()
        
        # Get available locales and their display names;
        available_locales = self.i18n.get_available_locales()
        display_names = [self.i18n.get_locale_display_name(locale) for locale in available_locales]
        
        lang_combo = ttk.Combobox(lang_frame, textvariable=self.lang_var, 
                                 values=display_names, 
                                 state='readonly', width=12)
        lang_combo.pack(side=tk.LEFT, padx=(5, 0))
        
        # Set current language display name;
        current_display_name = self.i18n.get_locale_display_name(self.i18n.current_locale)
        self.lang_var.set(current_display_name)
        
        lang_combo.bind('<<ComboboxSelected>>', self.on_language_change)
        
        # File selection area;
        file_frame = ttk.LabelFrame(main_frame, text=self.i18n.t('file_selection'), padding="10")
        file_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 20))
        file_frame.columnconfigure(1, weight=1)
        
        self.file_path_var = tk.StringVar()
        self.file_entry = ttk.Entry(file_frame, textvariable=self.file_path_var, state='readonly')
        self.file_entry.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), padx=(0, 10))
        
        self.browse_button = ttk.Button(file_frame, text=self.i18n.t('browse'), command=self.browse_file)
        self.browse_button.grid(row=0, column=2)
        
        # Parameters area;
        params_frame = ttk.LabelFrame(main_frame, text=self.i18n.t('detection_params'), padding="10")
        params_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 20))
        
        # Minimum duplicate length;
        ttk.Label(params_frame, text=self.i18n.t('min_duplicate_length')).grid(row=0, column=0, sticky=tk.W, padx=(0, 10))
        self.min_len_var = tk.IntVar(value=30)
        min_len_spinbox = ttk.Spinbox(params_frame, from_=10, to=100, width=10, 
                                     textvariable=self.min_len_var)
        min_len_spinbox.grid(row=0, column=1, sticky=tk.W)
        ttk.Label(params_frame, text=self.i18n.t('characters')).grid(row=0, column=2, sticky=tk.W, padx=(5, 0))
        
        # Action buttons;
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=5, column=0, columnspan=3, pady=(0, 20))
        
        self.process_button = ttk.Button(button_frame, text=self.i18n.t('start_detection'), 
                                        command=self.start_processing, style='Accent.TButton')
        self.process_button.pack(side=tk.LEFT, padx=(0, 10))
        
        self.reset_button = ttk.Button(button_frame, text=self.i18n.t('reset'), command=self.reset_form)
        self.reset_button.pack(side=tk.LEFT)
        
        # Progress bar;
        self.progress = ttk.Progressbar(main_frame, mode='indeterminate')
        self.progress.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(0, 10))
        
        # Status display area;
        status_frame = ttk.LabelFrame(main_frame, text=self.i18n.t('processing_status'), padding="10")
        status_frame.grid(row=7, column=0, columnspan=3, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 20))
        status_frame.columnconfigure(0, weight=1)
        status_frame.rowconfigure(0, weight=1)
        main_frame.rowconfigure(7, weight=1)
        
        # Status text box;
        self.status_text = tk.Text(status_frame, height=10, wrap=tk.WORD, state=tk.DISABLED)
        scrollbar = ttk.Scrollbar(status_frame, orient=tk.VERTICAL, command=self.status_text.yview)
        self.status_text.configure(yscrollcommand=scrollbar.set)
        
        self.status_text.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        
        # About button;
        about_button = ttk.Button(main_frame, text=self.i18n.t('about'), command=self.show_about)
        about_button.grid(row=8, column=2, sticky=tk.E)
        
        # Add initial prompt;
        self.log_message(self.i18n.t('select_file_prompt'))
    
    def on_language_change(self, event=None):
        """Handle language change;"""
        display_name = self.lang_var.get()
        locale_code = self.i18n.get_locale_from_display_name(display_name)
        if self.i18n.set_locale(locale_code):
            self.refresh_ui()
    
    def refresh_ui(self):
        """Refresh UI text after language change;"""
        self.root.title(self.i18n.t('window_title'))
        
        # Clear status and show language change message;
        self.status_text.config(state=tk.NORMAL)
        self.status_text.delete(1.0, tk.END)
        self.status_text.config(state=tk.DISABLED)
        
        # Recreate the entire UI with new language;
        for widget in self.root.winfo_children():
            widget.destroy()
        
        self.setup_ui()
    
    def browse_file(self):
        """Browse and select Word document;"""
        file_path = filedialog.askopenfilename(
            title=self.i18n.t('file_selection'),
            filetypes=[
                ("Word Documents", "*.docx"),
                ("All Files", "*.*")
            ]
        )
        if file_path:
            self.file_path_var.set(file_path)
            self.log_message(self.i18n.t('file_selected', filename=os.path.basename(file_path)))
    
    def reset_form(self):
        """Reset form;"""
        self.file_path_var.set("")
        self.min_len_var.set(30)
        self.status_text.config(state=tk.NORMAL)
        self.status_text.delete(1.0, tk.END)
        self.status_text.config(state=tk.DISABLED)
        self.log_message(self.i18n.t('form_reset'))
    
    def log_message(self, message):
        """Add message to status text box;"""
        self.status_text.config(state=tk.NORMAL)
        self.status_text.insert(tk.END, f"{message}\n")
        self.status_text.see(tk.END)
        self.status_text.config(state=tk.DISABLED)
        self.root.update_idletasks()
    
    def start_processing(self):
        """Start document processing;"""
        file_path = self.file_path_var.get().strip()
        
        if not file_path:
            messagebox.showerror(self.i18n.t('processing_failed_dialog'), self.i18n.t('error_no_file'))
            return
        
        if not os.path.exists(file_path):
            messagebox.showerror(self.i18n.t('processing_failed_dialog'), self.i18n.t('error_file_not_exist'))
            return
        
        if not file_path.lower().endswith('.docx'):
            messagebox.showerror(self.i18n.t('processing_failed_dialog'), self.i18n.t('error_invalid_format'))
            return
        
        # Disable buttons and start progress bar;
        self.process_button.config(state='disabled')
        self.browse_button.config(state='disabled')
        self.progress.start()
        
        # Process document in new thread;
        thread = threading.Thread(target=self.process_document, args=(file_path,))
        thread.daemon = True
        thread.start()
    
    def process_document(self, file_path):
        """Process document in background thread;"""
        try:
            self.log_message("="*50)
            self.log_message(self.i18n.t('processing_start', filename=os.path.basename(file_path)))
            self.log_message(self.i18n.t('min_length_setting', length=self.min_len_var.get()))
            self.log_message("="*50)
            
            # Use document processing logic;
            result = self.detect_and_add_comments(file_path)
            
            # Update UI in main thread;
            self.root.after(0, self.processing_complete, result)
            
        except Exception as e:
            error_msg = f"Error during processing: {str(e)}"
            self.root.after(0, self.processing_error, error_msg)
    
    def processing_complete(self, result):
        """Update UI after processing completion;"""
        self.progress.stop()
        self.process_button.config(state='normal')
        self.browse_button.config(state='normal')
        
        if result['success']:
            self.log_message(self.i18n.t('processing_complete'))
            self.log_message(self.i18n.t('output_file', filename=result['output_file']))
            self.log_message(self.i18n.t('comments_added', count=result['comment_count']))
            
            # Ask whether to open file folder;
            response = messagebox.askyesno(
                self.i18n.t('processing_complete'),
                self.i18n.t('processing_complete_dialog', 
                           count=result['comment_count'],
                           filename=os.path.basename(result['output_file']))
            )
            if response:
                self.open_file_location(result['output_file'])
        else:
            self.log_message(self.i18n.t('processing_failed', error=result['error']))
            messagebox.showerror(self.i18n.t('processing_failed_dialog'), result['error'])
    
    def processing_error(self, error_msg):
        """Update UI after processing error;"""
        self.progress.stop()
        self.process_button.config(state='normal')
        self.browse_button.config(state='normal')
        
        self.log_message(f"❌ {error_msg}")
        messagebox.showerror(self.i18n.t('processing_failed_dialog'), error_msg)
    
    def open_file_location(self, file_path):
        """Open file location in system file manager;"""
        try:
            if sys.platform == "win32":
                os.startfile(os.path.dirname(file_path))
            elif sys.platform == "darwin":
                os.system(f"open '{os.path.dirname(file_path)}'")
            else:
                os.system(f"xdg-open '{os.path.dirname(file_path)}'")
        except Exception as e:
            self.log_message(f"Cannot open file folder: {str(e)}")
    
    def show_about(self):
        """Show about dialog;"""
        messagebox.showinfo(self.i18n.t('about'), self.i18n.t('about_text'))
    
    # ============ Document Processing Core Logic ============
    
    def get_text_chunks(self, text, min_len=30):
        """Split text into chunks of specified length;"""
        return [text[i:i + min_len] for i in range(len(text) - min_len + 1)]

    def find_duplicate_info(self, texts, paragraphs, min_len=30):
        """Find duplicate information, return duplicate content, paragraph number and duplicate ID;"""
        results = []
        duplicate_groups = {}  # Track duplicate groups;
        duplicate_counter = 1  # Duplicate number counter;
        
        for i, base in enumerate(texts):
            if len(base) < min_len:
                results.append((None, None, None))
                continue
                
            chunks = self.get_text_chunks(base, min_len)
            found = None
            
            for j, other in enumerate(texts):
                if i == j:
                    continue
                for chunk in chunks:
                    if chunk in other:
                        # Check if this duplicate content already has a number;
                        chunk_key = chunk.strip()
                        if chunk_key not in duplicate_groups:
                            # New duplicate group, assign new number;
                            duplicate_groups[chunk_key] = duplicate_counter
                            duplicate_counter += 1
                        
                        # Return: duplicate content, source paragraph number, duplicate ID;
                        found = (chunk, j + 1, duplicate_groups[chunk_key])
                        break
                if found:
                    break
            results.append(found if found else (None, None, None))
        return results

    def create_comments_xml(self, comments):
        """Create Microsoft Office compatible comments.xml;"""
        # Define correct namespaces;
        w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        nsmap = {None: w_ns}
        
        # Create root element;
        root = etree.Element(f"{{{w_ns}}}comments", nsmap=nsmap)
        
        for comment_id, comment_text in comments:
            # Create comment element;
            comment_elem = etree.SubElement(root, f"{{{w_ns}}}comment")
            comment_elem.set(f"{{{w_ns}}}id", str(comment_id))
            comment_elem.set(f"{{{w_ns}}}author", "Word Duplicate Detector")
            comment_elem.set(f"{{{w_ns}}}date", "2025-07-10T00:00:00Z")
            comment_elem.set(f"{{{w_ns}}}initials", "WDD")
            
            # Create paragraph structure;
            p_elem = etree.SubElement(comment_elem, f"{{{w_ns}}}p")
            
            # Add paragraph properties;
            pPr = etree.SubElement(p_elem, f"{{{w_ns}}}pPr")
            pStyle = etree.SubElement(pPr, f"{{{w_ns}}}pStyle")
            pStyle.set(f"{{{w_ns}}}val", "CommentText")
            
            # Create run;
            r_elem = etree.SubElement(p_elem, f"{{{w_ns}}}r")
            
            # Add run properties;
            rPr = etree.SubElement(r_elem, f"{{{w_ns}}}rPr")
            rStyle = etree.SubElement(rPr, f"{{{w_ns}}}rStyle")
            rStyle.set(f"{{{w_ns}}}val", "CommentReference")
            
            # Create text element;
            t_elem = etree.SubElement(r_elem, f"{{{w_ns}}}t")
            if comment_text and (' ' in comment_text or comment_text.startswith(' ') or comment_text.endswith(' ')):
                t_elem.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t_elem.text = comment_text or ""
        
        # Generate XML string;
        xml_declaration = b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
        xml_content = etree.tostring(root, encoding='UTF-8', pretty_print=True)
        
        return xml_declaration + xml_content

    def add_comment_to_paragraph(self, paragraph, comment_text, comment_id, chunk_text):
        """Add comment to specific text in paragraph;"""
        # Get complete paragraph text;
        full_text = paragraph.text
        
        # Find position of text to annotate;
        chunk_index = full_text.find(chunk_text)
        if chunk_index == -1:
            return False
        
        # Clear paragraph content;
        paragraph.clear()
        
        # Rebuild paragraph content;
        before_text = full_text[:chunk_index]
        after_text = full_text[chunk_index + len(chunk_text):]
        
        # Add text before;
        if before_text:
            paragraph.add_run(before_text)
        
        # Add comment range start marker;
        p_elem = paragraph._element
        range_start = etree.Element(qn('w:commentRangeStart'))
        range_start.set(qn('w:id'), str(comment_id))
        p_elem.append(range_start)
        
        # Add text to annotate;
        target_run = paragraph.add_run(chunk_text)
        
        # Add comment range end marker;
        range_end = etree.Element(qn('w:commentRangeEnd'))
        range_end.set(qn('w:id'), str(comment_id))
        p_elem.append(range_end)
        
        # Add comment reference;
        comment_run_elem = etree.Element(qn('w:r'))
        comment_ref = etree.SubElement(comment_run_elem, qn('w:commentReference'))
        comment_ref.set(qn('w:id'), str(comment_id))
        p_elem.append(comment_run_elem)
        
        # Add text after;
        if after_text:
            paragraph.add_run(after_text)
        
        return True

    def inject_comments_into_docx(self, docx_path, output_path, comments):
        """Inject comments into DOCX file;"""
        if not comments:
            shutil.copy2(docx_path, output_path)
            return
        
        tmp_dir = f'tmp_docx_{os.getpid()}'
        
        try:
            # Clean up existing temporary directory;
            if os.path.exists(tmp_dir):
                shutil.rmtree(tmp_dir)
            
            # Extract DOCX file;
            with zipfile.ZipFile(docx_path, 'r') as zip_ref:
                zip_ref.extractall(tmp_dir)
            
            # Create comments.xml file;
            comments_dir = os.path.join(tmp_dir, 'word')
            os.makedirs(comments_dir, exist_ok=True)
            
            comments_xml_path = os.path.join(comments_dir, 'comments.xml')
            with open(comments_xml_path, 'wb') as f:
                f.write(self.create_comments_xml(comments))
            
            # Update [Content_Types].xml;
            self.update_content_types(tmp_dir)
            
            # Update document.xml.rels;
            self.update_document_rels(tmp_dir)
            
            # Repack DOCX file;
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED, compresslevel=6) as docx_zip:
                for root, dirs, files in os.walk(tmp_dir):
                    for file in files:
                        file_path = os.path.join(root, file)
                        arc_name = os.path.relpath(file_path, tmp_dir)
                        docx_zip.write(file_path, arc_name)
        
        finally:
            # Clean up temporary directory;
            if os.path.exists(tmp_dir):
                shutil.rmtree(tmp_dir)

    def update_content_types(self, tmp_dir):
        """Update Content_Types.xml file;"""
        content_types_path = os.path.join(tmp_dir, '[Content_Types].xml')
        
        if os.path.exists(content_types_path):
            tree = etree.parse(content_types_path)
            root = tree.getroot()
            
            # Check if comments content type already exists;
            ns = {'ct': 'http://schemas.openxmlformats.org/package/2006/content-types'}
            existing = root.find('.//ct:Override[@PartName="/word/comments.xml"]', namespaces=ns)
            
            if existing is None:
                # Add comments.xml content type;
                override = etree.SubElement(root, 
                    '{http://schemas.openxmlformats.org/package/2006/content-types}Override')
                override.set('PartName', '/word/comments.xml')
                override.set('ContentType', 
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml')
                
                # Save file;
                tree.write(content_types_path, xml_declaration=True, encoding='UTF-8', pretty_print=True)

    def update_document_rels(self, tmp_dir):
        """Update document.xml.rels file;"""
        rels_dir = os.path.join(tmp_dir, 'word', '_rels')
        rels_path = os.path.join(rels_dir, 'document.xml.rels')
        
        # Ensure _rels directory exists;
        os.makedirs(rels_dir, exist_ok=True)
        
        # Create basic structure if relationship file doesn't exist;
        if not os.path.exists(rels_path):
            nsmap = {None: 'http://schemas.openxmlformats.org/package/2006/relationships'}
            root = etree.Element('Relationships', nsmap=nsmap)
            tree = etree.ElementTree(root)
        else:
            tree = etree.parse(rels_path)
            root = tree.getroot()
        
        # Check if comments relationship already exists;
        ns = {'r': 'http://schemas.openxmlformats.org/package/2006/relationships'}
        existing_rel = None
        
        for rel in root.findall('r:Relationship', namespaces=ns):
            if rel.get('Type') == "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments":
                existing_rel = rel
                break
        
        if existing_rel is None:
            # Generate new relationship ID;
            existing_ids = [rel.get('Id', '') for rel in root.findall('r:Relationship', namespaces=ns)]
            max_id = 0
            for id_str in existing_ids:
                if id_str.startswith('rId') and id_str[3:].isdigit():
                    max_id = max(max_id, int(id_str[3:]))
            
            new_id = f"rId{max_id + 1}"
            
            # Add comments relationship;
            new_rel = etree.SubElement(root, 
                '{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')
            new_rel.set('Id', new_id)
            new_rel.set('Type', "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
            new_rel.set('Target', 'comments.xml')
            
            # Save file;
            tree.write(rels_path, xml_declaration=True, encoding='UTF-8', pretty_print=True)

    def detect_and_add_comments(self, input_path):
        """Detect duplicate content and add comments;"""
        try:
            self.log_message(self.i18n.t('loading_document'))
            
            # Read document;
            doc = Document(input_path)
            texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            paragraphs = [p for p in doc.paragraphs if p.text.strip()]
            
            self.log_message(self.i18n.t('document_loaded', count=len(texts)))
            self.log_message(self.i18n.t('detecting_duplicates'))
            
            # Detect duplicate content;
            min_len = self.min_len_var.get()
            duplicates = self.find_duplicate_info(texts, paragraphs, min_len)
            
            # Count duplicates;
            duplicate_count = sum(1 for chunk, origin, dup_id in duplicates if chunk is not None)
            self.log_message(self.i18n.t('detection_complete', count=duplicate_count))
            
            if duplicate_count == 0:
                # No duplicate content, copy file directly;
                base_name = os.path.splitext(input_path)[0]
                output_path = f"{base_name}_duplicate_check_result.docx"
                shutil.copy2(input_path, output_path)
                self.log_message(self.i18n.t('no_duplicates_found'))
                return {
                    'success': True,
                    'output_file': output_path,
                    'comment_count': 0
                }
            
            self.log_message(self.i18n.t('adding_comments'))
            
            # Prepare comment data;
            comments_to_add = []
            comment_id = 1
            
            # Process each paragraph;
            for i, (chunk, origin, dup_id) in enumerate(duplicates):
                if chunk is not None:
                    paragraph = paragraphs[i]
                    
                    # Generate comment text;
                    if self.i18n.current_locale == 'zh_Hans':
                        comment_text = f"重复#{dup_id}：与第{origin}段内容重复"
                    elif self.i18n.current_locale == 'zh_Hant':
                        comment_text = f"重複#{dup_id}：與第{origin}段內容重複"
                    else:
                        comment_text = f"Duplicate #{dup_id}: Repeats content from paragraph {origin}"
                    
                    log_info = self.i18n.t('comment_added', 
                                         paragraph=i+1, 
                                         group=dup_id, 
                                         origin=origin,
                                         preview=chunk[:25])
                    
                    # Add comment to paragraph;
                    success = self.add_comment_to_paragraph(paragraph, comment_text, comment_id, chunk)
                    if success:
                        comments_to_add.append((comment_id, comment_text))
                        comment_id += 1
                        self.log_message(log_info)
            
            # Save temporary file;
            base_name = os.path.splitext(input_path)[0]
            temp_path = f"{base_name}_temp.docx"
            output_path = f"{base_name}_with_comments.docx"
            
            self.log_message(self.i18n.t('saving_document'))
            doc.save(temp_path)
            
            # Inject comment XML;
            self.log_message(self.i18n.t('injecting_comments'))
            self.inject_comments_into_docx(temp_path, output_path, comments_to_add)
            
            # Clean up temporary file;
            if os.path.exists(temp_path):
                os.remove(temp_path)
            
            return {
                'success': True,
                'output_file': output_path,
                'comment_count': len(comments_to_add)
            }
            
        except Exception as e:
            return {
                'success': False,
                'error': str(e)
            }


def main():
    """Main function;"""
    # Create main window;
    root = tk.Tk()
    
    # Set Windows system DPI awareness;
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except:
        pass
    
    # Create application instance;
    app = DuplicateDetectorApp(root)
    
    # Run main loop;
    root.mainloop()


if __name__ == "__main__":
    main()
